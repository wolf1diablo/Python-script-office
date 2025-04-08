from googletrans import Translator
import docx
import os

def translate_docx(input_file, output_file, target_language='fr'):
    """
    Translates a .docx file from one language to another using googletrans.

    Args:
        input_file: Path to the input .docx file.
        output_file: Path to the output translated .docx file.
        target_language: The target language code (e.g., 'fr' for French).
    Raises:
        FileNotFoundError: If the input file does not exist.
        Exception: If there is an error during translation or file processing.
    """
    try:
        if not os.path.exists(input_file):
            raise FileNotFoundError(f"Input file '{input_file}' not found.")

        translator = Translator()
        doc = docx.Document(input_file)
        translated_doc = docx.Document()

        for paragraph in doc.paragraphs:
            try:
                translation = translator.translate(paragraph.text, dest=target_language)
                translated_doc.add_paragraph(translation.text)
            except Exception as e:
                print(f"Error translating paragraph: {paragraph.text.strip()}. Error: {e}")
                translated_doc.add_paragraph(paragraph.text) # Add original if translation fails

        translated_doc.save(output_file)
        print(f"File '{input_file}' translated to '{target_language}' and saved as '{output_file}'")

    except FileNotFoundError as e:
        print(f"Error: {e}")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")


if __name__ == "__main__":
    input_filename = "agreed.docx"  # Replace with your input .docx file
    output_filename = "the_output_fr.docx"  # Replace with your desired output file name
    target_lang = "fr"  # Replace with your desired target language code

    # Create a dummy docx file for testing
    #dummy_doc = docx.Document()
    #dummy_doc.add_paragraph("This is a test paragraph.")
    #dummy_doc.add_paragraph("Another paragraph to translate.")
    #dummy_doc.add_paragraph("This line has some special characters like éàçüöä. Let's see how it goes!")
    #dummy_doc.save(input_filename)


    translate_docx(input_filename, output_filename, target_lang)

    ## Example changing target language to Spanish
    #output_filename_es = "output_es.docx"
    #target_lang_es = "es"
    #translate_docx(input_filename, output_filename_es, target_lang_es)

    # Clean up the dummy file (Optional)
    #os.remove(input_filename)